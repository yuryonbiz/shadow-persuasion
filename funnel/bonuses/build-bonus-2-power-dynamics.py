#!/usr/bin/env python3
"""
Bonus #2: The Power Dynamics Cheatsheet

A single-page reference designed to be screenshotted to a phone and
pulled up before any high-stakes conversation.

Seven sections:
  1. Pre-Flight Check (what to settle before you walk in)
  2. First 30 Seconds (opening moves)
  3. Who Has The Power (real-time read)
  4. Power Moves (deploy)
  5. Power Leaks (avoid)
  6. When You're Losing (recovery)
  7. Exit Well (ending the conversation)

Generates:
  - cover-bonus-2.png (cream cover image)
  - bonus-2-power-dynamics-cheatsheet.docx
    Page 1: cover
    Page 2: the cheatsheet itself (screenshot-friendly)
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "public" / "logo.png"
FONT_TTF = REPO / "funnel" / "assets" / "SpecialElite-Regular.ttf"
COVER_OUT = REPO / "funnel" / "assets" / "cover-bonus-2.png"
DOCX_OUT = REPO / "funnel" / "exports" / "bonus-2-power-dynamics-cheatsheet.docx"

BODY_FONT = "Georgia"
TITLE_FONT = "Georgia"

# Brand colors (RGB tuples for Pillow, RGBColor for docx)
CREAM = (0xF4, 0xEC, 0xD8)
DARK = (0x1A, 0x1A, 0x1A)
BROWN = (0x5C, 0x3A, 0x1E)
GOLD = (0xD4, 0xA0, 0x17)
RED = (0x8B, 0x3D, 0x3D)

DARK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
BROWN_RGB = RGBColor(0x5C, 0x3A, 0x1E)
GOLD_RGB = RGBColor(0xD4, 0xA0, 0x17)
RED_RGB = RGBColor(0x8B, 0x3D, 0x3D)
GREEN_RGB = RGBColor(0x3D, 0x8B, 0x5E)
GRAY_RGB = RGBColor(0x6B, 0x5B, 0x3E)


# ========== COVER ==========

def build_cover():
    WIDTH, HEIGHT = 1800, 2700
    img = Image.new("RGB", (WIDTH, HEIGHT), CREAM)
    draw = ImageDraw.Draw(img)

    label_font = ImageFont.truetype(str(FONT_TTF), 40)
    title_font = ImageFont.truetype(str(FONT_TTF), 100)
    subtitle_font = ImageFont.truetype(str(FONT_TTF), 58)
    author_font = ImageFont.truetype(str(FONT_TTF), 44)

    # Logo
    logo = Image.open(str(LOGO)).convert("RGBA")
    scale = 900 / logo.width
    logo_resized = logo.resize((int(logo.width * scale), int(logo.height * scale)), Image.LANCZOS)
    img.paste(logo_resized, ((WIDTH - logo_resized.width) // 2, 360), logo_resized)

    # Bonus label
    label_text = "// FREE BONUS  02 //"
    bbox = draw.textbbox((0, 0), label_text, font=label_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, 880), label_text, fill=BROWN, font=label_font)

    # Title
    title_text = "THE POWER"
    bbox = draw.textbbox((0, 0), title_text, font=title_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, 1020), title_text, fill=DARK, font=title_font)

    title_text2 = "DYNAMICS"
    bbox = draw.textbbox((0, 0), title_text2, font=title_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, 1150), title_text2, fill=DARK, font=title_font)

    title_text3 = "CHEATSHEET"
    bbox = draw.textbbox((0, 0), title_text3, font=title_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, 1280), title_text3, fill=DARK, font=title_font)

    # Gold rule
    draw.rectangle([(WIDTH // 2 - 180, 1450), (WIDTH // 2 + 180, 1453)], fill=GOLD)

    # Subtitle
    subtitle_lines = [
        "Screenshot This.",
        "Pull It Up Before",
        "Any Conversation",
        "That Matters.",
    ]
    sub_y = 1520
    for line in subtitle_lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, sub_y), line, fill=BROWN, font=subtitle_font)
        sub_y += 90

    # Author
    author_text = "NATE HARLAN"
    bbox = draw.textbbox((0, 0), author_text, font=author_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, HEIGHT - 280), author_text, fill=DARK, font=author_font)

    # Border
    draw.rectangle([(60, 60), (WIDTH - 60, HEIGHT - 60)], outline=BROWN, width=2)

    COVER_OUT.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(COVER_OUT), "PNG", dpi=(300, 300))
    print(f"Cover saved: {COVER_OUT}")


# ========== DOCX HELPERS ==========

def add_page_number(paragraph):
    run = paragraph.add_run()
    for t, val in [("w:fldChar", "begin")]:
        el = OxmlElement(t)
        el.set(qn("w:fldCharType"), val)
        run._r.append(el)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    run._r.append(it)
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    run._r.append(el)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BROWN_RGB


def set_page_background_cream(doc):
    """Set the page background to cream (#F4ECD8). Shows on screen and exports to PDF."""
    background = OxmlElement("w:background")
    background.set(qn("w:color"), "F4ECD8")
    doc.element.insert(0, background)
    # Also need displayBackgroundShape in settings for Word to render it
    settings = doc.settings.element
    display_bg = OxmlElement("w:displayBackgroundShape")
    settings.append(display_bg)


# ========== CHEATSHEET CONTENT ==========

SECTIONS = [
    {
        "title": "PRE-FLIGHT CHECK",
        "subtitle": "Before you walk in.",
        "items": [
            "Know your BATNA. What happens if this goes badly.",
            "Pick ONE tactic to deploy. Not three.",
            "Your walkaway number stays in your head. Never spoken.",
        ],
        "label_color": BROWN_RGB,
    },
    {
        "title": "FIRST 30 SECONDS",
        "subtitle": "The window that determines everything.",
        "items": [
            "Match their energy. Do not exceed it.",
            "Skip the \"thanks for making time.\" Open with a specific observation.",
            "Silence is okay. Filling it is not.",
        ],
        "label_color": BROWN_RGB,
    },
    {
        "title": "WHO HAS THE POWER",
        "subtitle": "A two-second read, any moment in the conversation.",
        "items": [
            "YOU hold it if: they're asking more questions than you are.",
            "YOU hold it if: silences don't make you uncomfortable.",
            "THEY hold it if: you're asking more questions than they are.",
            "THEY hold it if: you're filling every pause.",
        ],
        "label_color": GOLD_RGB,
    },
    {
        "title": "POWER MOVES",
        "subtitle": "Deploy when you feel the dynamic shifting toward you.",
        "items": [
            "Four seconds of silence after their statement.",
            "Ask a \"how\" question when cornered.",
            "Let them name the next step.",
            "Short emails beat long ones. Twelve words is often enough.",
        ],
        "label_color": GREEN_RGB,
    },
    {
        "title": "POWER LEAKS",
        "subtitle": "Each of these cedes leverage you did not mean to give.",
        "items": [
            "Over-explaining when they did not ask.",
            "Apologizing for taking their time.",
            "Laughing at every joke to build rapport.",
            "Using \"just\" and \"actually\" as softeners.",
            "Answering questions you were not asked.",
        ],
        "label_color": RED_RGB,
    },
    {
        "title": "WHEN YOU'RE LOSING",
        "subtitle": "Recovery moves. Use in order.",
        "items": [
            "Stop. Silence for four seconds.",
            "Ask a question. Any question at all.",
            "Restate your core point verbatim. Do not re-argue it.",
            "Worst case: \"Let me think. I'll follow up Monday.\"",
        ],
        "label_color": RED_RGB,
    },
    {
        "title": "EXIT WELL",
        "subtitle": "The last 30 seconds matter as much as the first.",
        "items": [
            "Do not celebrate a yes. Move to an unrelated topic.",
            "Do not summarize what was agreed.",
            "Twelve-word email within twenty-four hours.",
            "Choose the emotional state you leave them in. That is the memory.",
        ],
        "label_color": BROWN_RGB,
    },
]


# ========== DOCX BUILD ==========

def build_docx():
    doc = Document()
    set_page_background_cream(doc)

    # ==== SECTION 1: COVER (full-bleed) ====
    cover_section = doc.sections[0]
    cover_section.page_width = Inches(6)
    cover_section.page_height = Inches(9)
    cover_section.top_margin = Inches(0)
    cover_section.bottom_margin = Inches(0)
    cover_section.left_margin = Inches(0)
    cover_section.right_margin = Inches(0)
    cover_section.header_distance = Inches(0)
    cover_section.footer_distance = Inches(0)

    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1.0
    run = cover_p.add_run()
    run.add_picture(str(COVER_OUT), width=Inches(6), height=Inches(9))

    # ==== SECTION 2: THE CHEATSHEET (single page) ====
    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body.page_width = Inches(6)
    body.page_height = Inches(9)
    body.top_margin = Inches(0.6)
    body.bottom_margin = Inches(0.5)
    body.left_margin = Inches(0.55)
    body.right_margin = Inches(0.55)
    body.header_distance = Inches(0.2)
    body.footer_distance = Inches(0.3)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False

    # Cheatsheet has no header (one-page layout; logo lives in title area)
    # Footer: page number (will just say "2" on this single content page)
    f = body.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = f.add_run("SHADOWPERSUASION.COM")
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = BROWN_RGB

    # === Title block ===
    # Logo at top right (small inline)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run()
    run.add_picture(str(LOGO), width=Inches(1.2))

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(6)
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run("POWER DYNAMICS CHEATSHEET")
    r.font.name = TITLE_FONT
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(8)
    r = s.add_run("Screenshot this. Check it before any conversation that matters.")
    r.font.name = TITLE_FONT
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = BROWN_RGB

    # Separator rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(6)
    rr = rule.add_run("_" * 40)
    rr.font.name = BODY_FONT
    rr.font.size = Pt(8)
    rr.font.color.rgb = GOLD_RGB

    # === The seven sections ===
    for idx, section in enumerate(SECTIONS):
        # Section header
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(1)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(section["title"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = section["label_color"]

        # Section subtitle
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(3)
        sub.paragraph_format.keep_with_next = True
        r = sub.add_run(section["subtitle"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(8.5)
        r.italic = True
        r.font.color.rgb = GRAY_RGB

        # Items
        for i, item in enumerate(section["items"]):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.keep_with_next = (i < len(section["items"]) - 1)
            # Use a plain bullet
            r1 = p.add_run("·  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(9.5)
            r1.bold = True
            r1.font.color.rgb = section["label_color"]
            r2 = p.add_run(item)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = DARK_RGB

    # Final brand mark at bottom
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(12)
    r = foot.add_run("From Shadow Persuasion by Nate Harlan")
    r.font.name = TITLE_FONT
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = BROWN_RGB

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"Docx saved: {DOCX_OUT}")
    print(f"Size: {DOCX_OUT.stat().st_size} bytes")


if __name__ == "__main__":
    build_cover()
    build_docx()
