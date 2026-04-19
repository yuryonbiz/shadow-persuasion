#!/usr/bin/env python3
"""
Build a professionally-formatted Word document for Shadow Persuasion.

Features:
- Cover page with centered large logo + title + subtitle + author
- Table of Contents on its own page
- Each chapter starts on a new page
- Header on all content pages with small logo right-aligned
- Footer on all content pages with page number centered
- Body font: EB Garamond (fallback Georgia) for readability
- Heading font: Special_Elite (fallback Courier New) for brand match
- Italic "Field Note" paragraphs styled distinctively
- Zero horizontal rules anywhere
- No duplicate title block after TOC
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import re

REPO = Path(__file__).resolve().parent.parent
MANUSCRIPT = REPO / "funnel" / "shadow-persuasion-full-manuscript.md"
LOGO = REPO / "public" / "logo.png"
OUT = REPO / "funnel" / "exports" / "shadow-persuasion.docx"

BODY_FONT = "EB Garamond"
HEADING_FONT = "Courier New"  # falls back if Special_Elite missing
TITLE_FONT = "Georgia"

# ----- Helpers -----

def add_page_number(paragraph):
    """Insert a PAGE field into a paragraph (live page number)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x5C, 0x3A, 0x1E)


def set_page_break_before(paragraph):
    """Force paragraph to start on new page."""
    pPr = paragraph._p.get_or_add_pPr()
    pageBreakBefore = OxmlElement("w:pageBreakBefore")
    pPr.append(pageBreakBefore)


def insert_page_break(paragraph):
    """Insert a page break as a run inside paragraph."""
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_toc(document):
    """Add a Table of Contents field. Word will populate on open (user hits F9 to refresh)."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and select Update Field to populate the table of contents."
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)


def style_paragraph(p, font=BODY_FONT, size=11, italic=False, bold=False, align=None, color=None, space_after=6, first_indent=None, line=1.35):
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if first_indent is not None:
        pf.first_line_indent = first_indent
    for run in p.runs:
        run.font.name = font
        run.font.size = Pt(size)
        run.italic = italic
        run.bold = bold
        if color:
            run.font.color.rgb = color


# ----- Markdown parsing (simple, targeted to this book's format) -----

INLINE_PATTERNS = [
    (re.compile(r"\*\*(.+?)\*\*"), "bold"),   # **bold**
    (re.compile(r"\*(.+?)\*"), "italic"),     # *italic*
    (re.compile(r"_(.+?)_"), "italic"),       # _italic_
]


def add_inline_formatted(paragraph, text, font=BODY_FONT, size=11, base_italic=False, color=None):
    """Add text to paragraph, interpreting simple markdown inline syntax."""
    # Tokenize: split by **bold** and *italic* preserving content
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for tok in tokens:
        if not tok:
            continue
        bold = italic = False
        content = tok
        if tok.startswith("**") and tok.endswith("**"):
            bold = True
            content = tok[2:-2]
        elif tok.startswith("*") and tok.endswith("*") and len(tok) >= 2:
            italic = True
            content = tok[1:-1]
        run = paragraph.add_run(content)
        run.font.name = font
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic or base_italic
        if color:
            run.font.color.rgb = color


# ----- Build the document -----

def build_document():
    doc = Document()

    # Page setup for all sections — 6x9 book trim with comfortable margins.
    for section in doc.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    section = doc.sections[0]
    # Disable header/footer on first page (cover)
    section.different_first_page_header_footer = True

    # =========== COVER PAGE ===========
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Some vertical space
    cover.paragraph_format.space_before = Pt(120)
    r = cover.add_run()
    r.add_picture(str(LOGO), width=Inches(3.5))

    # Title (spacer + big title)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(60)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("SHADOW PERSUASION")
    tr.font.name = TITLE_FONT
    tr.font.size = Pt(28)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    title.paragraph_format.space_after = Pt(6)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("The 47 Counterintuitive Conversation Tactics\nThat Make People Say Yes Without Realizing Why")
    sr.font.name = TITLE_FONT
    sr.font.size = Pt(14)
    sr.italic = True
    sr.font.color.rgb = RGBColor(0x5C, 0x3A, 0x1E)
    sub.paragraph_format.space_after = Pt(48)

    # Author
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run("NATE HARLAN")
    ar.font.name = TITLE_FONT
    ar.font.size = Pt(14)
    ar.bold = True
    ar.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Page break after cover
    pb = doc.add_paragraph()
    insert_page_break(pb)

    # =========== HEADER / FOOTER for non-cover pages ===========
    # Header right-aligned logo
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = header_para.add_run()
    hrun.add_picture(str(LOGO), width=Inches(1.0))

    # Footer centered page number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)

    # =========== COPYRIGHT PAGE ===========
    copyright_lines = [
        ("Copyright © 2026 Shadow Persuasion LLC.", False),
        ("All rights reserved.", False),
        ("", False),
        ("No part of this book may be reproduced, stored in a retrieval system, or transmitted in any form or by any means, electronic, mechanical, photocopying, recording, or otherwise, without the prior written permission of the copyright holder, except in the case of brief quotations in critical articles or reviews.", True),
        ("", False),
        ("Case studies in this book are composites based on real coaching engagements, with names and identifying details modified to protect privacy. Tactical content is grounded in published research where cited.", True),
        ("", False),
        ("This book is not a substitute for professional legal, medical, financial, or therapeutic advice. Nothing in this book is a guarantee of any specific outcome.", True),
        ("", False),
        ("First edition.", False),
        ("", False),
        ("Shadow Persuasion LLC", False),
        ("Boulder, Colorado", False),
    ]

    for text, wrap in copyright_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            r = p.add_run(text)
            r.font.name = BODY_FONT
            r.font.size = Pt(9 if wrap else 10)
            r.font.color.rgb = RGBColor(0x3B, 0x2E, 0x1A)
        p.paragraph_format.space_after = Pt(4)

    # Page break
    pb = doc.add_paragraph()
    insert_page_break(pb)

    # =========== TABLE OF CONTENTS ===========
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = toc_title.add_run("CONTENTS")
    tr.font.name = TITLE_FONT
    tr.font.size = Pt(22)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    toc_title.paragraph_format.space_after = Pt(24)

    add_toc(doc)

    pb = doc.add_paragraph()
    insert_page_break(pb)

    # =========== BODY CONTENT (from markdown) ===========
    # Parse the full manuscript. Skip front matter and introduction title metadata
    # since we added those above. Begin with Introduction heading.

    source = MANUSCRIPT.read_text()

    # Strip all horizontal rules (--- and ***) anywhere they appear as a full line
    lines = source.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if stripped in ("---", "***", "----", "----"):
            continue
        if re.fullmatch(r"-{3,}", stripped) or re.fullmatch(r"\*{3,}", stripped):
            continue
        cleaned.append(line)

    text = "\n".join(cleaned)

    # Remove the front-matter block entirely — we built the cover/copyright manually above.
    # Front matter ends before the "# Introduction" heading.
    intro_idx = text.find("# Introduction")
    if intro_idx > 0:
        text = text[intro_idx:]

    # Now parse block by block
    paragraphs = text.split("\n\n")

    in_intro = False
    first_chapter_seen = False

    for block in paragraphs:
        block = block.strip()
        if not block:
            continue

        # Detect heading level
        if block.startswith("# "):
            heading_text = block[2:].strip()

            # Render as Chapter Title style
            p = doc.add_paragraph()
            # First chapter also triggers page break (consistent with copyright/TOC flow)
            if first_chapter_seen:
                set_page_break_before(p)
            first_chapter_seen = True

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(12)
            p.style = doc.styles["Heading 1"]
            # override the style for our look
            r = p.add_run(heading_text.upper())
            r.font.name = TITLE_FONT
            r.font.size = Pt(20)
            r.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        elif block.startswith("## "):
            heading_text = block[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(24)
            p.style = doc.styles["Heading 2"]
            r = p.add_run(heading_text)
            r.font.name = TITLE_FONT
            r.font.size = Pt(14)
            r.italic = True
            r.font.color.rgb = RGBColor(0x5C, 0x3A, 0x1E)

        elif block.startswith("### "):
            heading_text = block[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
            p.style = doc.styles["Heading 3"]
            r = p.add_run(heading_text)
            r.font.name = TITLE_FONT
            r.font.size = Pt(13)
            r.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        elif block.startswith("- ") or block.startswith("* "):
            # Bullet list — split on lines
            for line in block.splitlines():
                line = line.strip()
                if not line:
                    continue
                content = line.lstrip("-* ").strip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = 1.3
                p.paragraph_format.space_after = Pt(4)
                add_inline_formatted(p, content, font=BODY_FONT, size=11)

        else:
            # Normal paragraph (may have **bold**, *italic*, or be a full-italic Field Note)
            # Detect Field Note: entire block wrapped in *...*
            stripped = block.strip()
            if stripped.startswith("*") and stripped.endswith("*") and stripped.count("\n") < 3 and not stripped.startswith("**"):
                # Italic Field Note paragraph
                inner = stripped[1:-1].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(18)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                r = p.add_run(inner)
                r.font.name = BODY_FONT
                r.font.size = Pt(11)
                r.italic = True
                r.font.color.rgb = RGBColor(0x5C, 0x3A, 0x1E)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.35
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Inches(0.25)
                add_inline_formatted(p, block.replace("\n", " "), font=BODY_FONT, size=11)

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    build_document()
