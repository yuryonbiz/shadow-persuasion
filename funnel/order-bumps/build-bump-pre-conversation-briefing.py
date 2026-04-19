#!/usr/bin/env python3
"""
Order Bump: The Pre-Conversation Briefing.

A 10-minute fillable worksheet + companion guide you fill out before any
high-stakes conversation. Solves the three problems the book creates:
- Can't remember 47 tactics in real-time
- Don't know which tactic fits MY specific situation
- Nerves make me sound wrong even when I know what to say

Price: $17 order bump at checkout.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "public" / "logo.png"
FONT_TTF = REPO / "funnel" / "assets" / "SpecialElite-Regular.ttf"
COVER_OUT = REPO / "funnel" / "assets" / "cover-bump-briefing.png"
DOCX_OUT = REPO / "funnel" / "exports" / "bump-pre-conversation-briefing.docx"

BODY_FONT = "Georgia"
TITLE_FONT = "Georgia"

CREAM = (0xF4, 0xEC, 0xD8)
DARK = (0x1A, 0x1A, 0x1A)
BROWN = (0x5C, 0x3A, 0x1E)
GOLD = (0xD4, 0xA0, 0x17)

DARK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
BROWN_RGB = RGBColor(0x5C, 0x3A, 0x1E)
GOLD_RGB = RGBColor(0xD4, 0xA0, 0x17)
GRAY_RGB = RGBColor(0x6B, 0x5B, 0x3E)


# ========== COVER ==========

def build_cover():
    WIDTH, HEIGHT = 1800, 2700
    img = Image.new("RGB", (WIDTH, HEIGHT), CREAM)
    draw = ImageDraw.Draw(img)

    label_font = ImageFont.truetype(str(FONT_TTF), 40)
    title_font = ImageFont.truetype(str(FONT_TTF), 110)
    subtitle_font = ImageFont.truetype(str(FONT_TTF), 58)
    stamp_font = ImageFont.truetype(str(FONT_TTF), 36)
    author_font = ImageFont.truetype(str(FONT_TTF), 52)

    # Logo
    logo = Image.open(str(LOGO)).convert("RGBA")
    scale = 900 / logo.width
    logo_resized = logo.resize((int(logo.width * scale), int(logo.height * scale)), Image.LANCZOS)
    logo_x = (WIDTH - logo_resized.width) // 2
    logo_y = 380
    img.paste(logo_resized, (logo_x, logo_y), logo_resized)

    # Label
    label_text = "// THE BRIEFING //"
    bbox = draw.textbbox((0, 0), label_text, font=label_font)
    label_y = logo_y + logo_resized.height + 80
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, label_y), label_text, fill=BROWN, font=label_font)

    # Title — 3 lines
    title_y = label_y + 120
    title_line_spacing = 135
    title_lines = ["THE PRE-", "CONVERSATION", "BRIEFING"]
    for i, line in enumerate(title_lines):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, title_y + i * title_line_spacing),
                  line, fill=DARK, font=title_font)

    # Gold rule
    rule_y = title_y + (len(title_lines) * title_line_spacing) + 60
    draw.rectangle([(WIDTH // 2 - 200, rule_y), (WIDTH // 2 + 200, rule_y + 3)], fill=GOLD)

    # Subtitle — 3 lines
    subtitle_lines = [
        "The 10-Minute System",
        "You Fill Out Before Any",
        "Conversation That Matters",
    ]
    sub_y = rule_y + 80
    for line in subtitle_lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, sub_y), line, fill=BROWN, font=subtitle_font)
        sub_y += 90

    # Stamp/badge — rotated-looking
    stamp_w, stamp_h = 520, 110
    stamp_x = (WIDTH - stamp_w) // 2
    stamp_y = sub_y + 90
    draw.rectangle([(stamp_x, stamp_y), (stamp_x + stamp_w, stamp_y + stamp_h)], outline=DARK, width=4)
    stamp_text = "FILL // REVIEW // WIN"
    bbox = draw.textbbox((0, 0), stamp_text, font=stamp_font)
    sx = stamp_x + (stamp_w - (bbox[2] - bbox[0])) // 2
    sy = stamp_y + (stamp_h - (bbox[3] - bbox[1])) // 2 - 8
    draw.text((sx, sy), stamp_text, fill=DARK, font=stamp_font)

    # Author
    author_text = "NATE HARLAN"
    bbox = draw.textbbox((0, 0), author_text, font=author_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, HEIGHT - 280), author_text, fill=DARK, font=author_font)

    # Border
    draw.rectangle([(60, 60), (WIDTH - 60, HEIGHT - 60)], outline=BROWN, width=2)

    COVER_OUT.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(COVER_OUT), "PNG", dpi=(300, 300))
    print(f"Cover saved: {COVER_OUT}")


# ========== DOCX HELPERS ==========

def add_page_number(paragraph):
    run = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "begin")
    run._r.append(el)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    run._r.append(it)
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    run._r.append(el)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BROWN_RGB


def set_page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


def heading(doc, text, size=18, center=True, space_before=14, space_after=8, color=None):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = TITLE_FONT
    r.font.size = Pt(size)
    r.bold = True
    r.font.color.rgb = color if color else DARK_RGB
    return p


def body(doc, text, size=11, first_indent=0.25, space_after=8, justify=True, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    if italic:
        r.italic = True
    r.font.color.rgb = color if color else DARK_RGB
    return p


def worksheet_prompt(doc, num, prompt_text, guidance=None, lines=3):
    """Big prompt with underline-space for writing."""
    # Prompt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(f"{num}.  ")
    r1.font.name = TITLE_FONT
    r1.font.size = Pt(12)
    r1.bold = True
    r1.font.color.rgb = GOLD_RGB
    r2 = p.add_run(prompt_text)
    r2.font.name = TITLE_FONT
    r2.font.size = Pt(12)
    r2.bold = True
    r2.font.color.rgb = DARK_RGB

    if guidance:
        g = doc.add_paragraph()
        g.paragraph_format.space_after = Pt(4)
        g.paragraph_format.left_indent = Inches(0.25)
        g.paragraph_format.keep_with_next = True
        r = g.add_run(guidance)
        r.font.name = BODY_FONT
        r.font.size = Pt(9.5)
        r.italic = True
        r.font.color.rgb = GRAY_RGB

    # Fillable lines
    for _ in range(lines):
        ln = doc.add_paragraph()
        ln.paragraph_format.space_after = Pt(2)
        ln.paragraph_format.line_spacing = 1.3
        r = ln.add_run(" " + "_" * 90)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY_RGB


def rule(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("\u2500 \u2500 \u2500")
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.color.rgb = GOLD_RGB


# ========== DOCX BUILD ==========

def build_docx():
    doc = Document()

    # ==== SECTION 1: COVER (zero-margin) ====
    cover_section = doc.sections[0]
    cover_section.page_width = Inches(6)
    cover_section.page_height = Inches(9)
    cover_section.top_margin = Inches(0)
    cover_section.bottom_margin = Inches(0)
    cover_section.left_margin = Inches(0)
    cover_section.right_margin = Inches(0)
    cover_section.header_distance = Inches(0)
    cover_section.footer_distance = Inches(0)

    cover_p = doc.add_paragraph()
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1.0
    run = cover_p.add_run()
    run.add_picture(str(COVER_OUT), width=Inches(6), height=Inches(9))

    # ==== SECTION 2: BODY ====
    b = doc.add_section(WD_SECTION_START.NEW_PAGE)
    b.page_width = Inches(6)
    b.page_height = Inches(9)
    b.top_margin = Inches(0.9)
    b.bottom_margin = Inches(0.8)
    b.left_margin = Inches(0.7)
    b.right_margin = Inches(0.7)
    b.header_distance = Inches(0.3)
    b.footer_distance = Inches(0.4)
    b.header.is_linked_to_previous = False
    b.footer.is_linked_to_previous = False

    # Header logo right
    h = b.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = h.add_run()
    hrun.add_picture(str(LOGO), width=Inches(0.9))

    # Footer page #
    f = b.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(f)

    # ==== TITLE PAGE ====
    for _ in range(2):
        doc.add_paragraph()

    heading(doc, "THE PRE-CONVERSATION BRIEFING", size=22, space_before=0, space_after=6)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(20)
    r = st.add_run("The 10-Minute System You Fill Out Before Any Conversation That Matters")
    r.font.name = TITLE_FONT
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = BROWN_RGB

    by = doc.add_paragraph()
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = by.add_run("by Nate Harlan")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_RGB

    # ==== PART 1: HOW TO USE ====
    heading(doc, "HOW TO USE THIS BRIEFING", size=18, space_before=36)

    intro = [
        "You read Shadow Persuasion. You now know 47 tactics exist. Here's the problem I didn't fully solve for you in the book: when the real conversation starts Wednesday morning, you won't remember which tactic fits, you'll be nervous enough that your own detector fires, and you'll come out saying the thing you promised yourself you wouldn't say.",
        "That's not a knowledge problem. It's a preparation problem. This briefing solves it.",
        "The briefing is a ten-minute form you fill out before any conversation that matters. It forces you to pick 3 of the 47 tactics that fit your specific situation, draft your opening line, anticipate the pushback, pre-write your responses, run a 90-second nerves routine, and walk in with all of it loaded into short-term memory. Everyone I work with privately uses a version of this before every high-stakes call.",
        "The payoff compounds. Every briefing you fill out makes the next one faster, because you start recognizing your own patterns. Three months in, the briefing takes five minutes. Six months in, you run it in your head while you walk to the meeting.",
    ]
    for para in intro:
        body(doc, para)

    heading(doc, "The 90 / 30 / 10 Rhythm", size=14, space_before=18, center=False)

    body(doc,
         "Use this briefing on a specific timing rhythm for best results. Block time intentionally. The calendar is the first tactic.",
         italic=True, color=BROWN_RGB, first_indent=0)

    cycle = [
        ("90 minutes before the conversation:",
         "Fill out the worksheet. Draft your opening line. Pre-write your responses to the two objections you most expect. Pick your three tactics. Identify your exit move."),
        ("30 minutes before:",
         "Run the nerves routine on the next page. Breath, body, reframe. Do not look at your phone during this window."),
        ("10 minutes before:",
         "Re-read the worksheet one time. Close it. Walk into the conversation with it in short-term memory, not in front of you."),
        ("Within 2 hours after:",
         "Fill out the debrief page. Even if it went well. Especially if it didn't. The debrief is where the next briefing gets sharper."),
    ]
    for head, text in cycle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.15)
        r1 = p.add_run(head + "  ")
        r1.font.name = BODY_FONT
        r1.font.size = Pt(11)
        r1.bold = True
        r1.font.color.rgb = GOLD_RGB
        r2 = p.add_run(text)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK_RGB

    # ==== PART 2: NERVES ROUTINE ====
    heading(doc, "THE 90-SECOND NERVES ROUTINE", size=18, space_before=28)
    set_page_break_before(doc.paragraphs[-1])

    body(doc, "The book taught you to appear calm and low-status. That only works if you actually are calm. Nerves leak. A wobbly voice, a rushed pace, or over-explaining all fire your OWN detector-giveaways and the other side reads it. This 90-second routine resets your nervous system. It is not optional before any high-stakes conversation.")

    heading(doc, "STEP 1 — The 4-7-8 Breath (45 seconds)", size=13, space_before=14, center=False, color=GOLD_RGB)
    body(doc, "Inhale through the nose for 4 seconds. Hold for 7. Exhale through the mouth slowly for 8. Three full cycles, 57 seconds total. This lowers heart rate and activates the parasympathetic nervous system. It's the fastest physiological intervention you have. Works every time.", first_indent=0)

    heading(doc, "STEP 2 — The Physical Reset (30 seconds)", size=13, space_before=14, center=False, color=GOLD_RGB)
    body(doc, "Roll your shoulders back and down. Unclench your jaw. Loosen your hands (nervous people clench them without noticing). Stand up if you're sitting, sit down if you're standing \u2014 the state shift breaks the anxiety loop. Smile for three seconds. Not because you feel like it. Because your facial muscles signal your brain that the situation is safe.", first_indent=0)

    heading(doc, "STEP 3 — The Reframe (15 seconds)", size=13, space_before=14, center=False, color=GOLD_RGB)
    body(doc, "Say this in your head, slowly: \u201cThis is a conversation. I've had thousands of conversations. I will have thousands more. The outcome of this one does not define me. My bottom line is [fill in your worksheet number]. If we don't reach it, I walk away and come back later. That's it.\u201d", first_indent=0)

    body(doc, "The reframe works because it re-installs proportion. Nerves come from treating a conversation as existential. The reframe reminds your brain it isn't.", italic=True, first_indent=0, color=BROWN_RGB)

    # ==== PART 3: THE BRIEFING WORKSHEET ====
    heading(doc, "THE BRIEFING WORKSHEET", size=20, space_before=28)
    set_page_break_before(doc.paragraphs[-1])

    body(doc, "Fill out the following in under ten minutes. Be honest. The worksheet only works if you actually write answers \u2014 not if you think you could if you wanted to. Handwritten works better than typed for retention, but either is fine.", italic=True, color=BROWN_RGB)

    # Meta
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("DATE:  " + "_" * 20 + "     TIME OF CONVERSATION:  " + "_" * 15)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_RGB

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("OTHER PERSON(S) IN THE CONVERSATION:  " + "_" * 40)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_RGB

    rule(doc)

    worksheet_prompt(doc, "01",
                     "What am I walking into?",
                     "One sentence. Name the conversation in plain English. \u201cSalary review with my manager.\u201d \u201cCustody mediation with my ex.\u201d \u201cReconnection call with my sister after 2 years.\u201d Don't editorialize. Name it.",
                     lines=3)

    worksheet_prompt(doc, "02",
                     "What outcome do I want?",
                     "The specific result. Not \u201cthe meeting goes well.\u201d \u201c$140K base salary starting October.\u201d \u201cWe agree on every-other-weekend custody.\u201d \u201cWe schedule a second conversation within a week.\u201d",
                     lines=3)

    worksheet_prompt(doc, "03",
                     "What does the OTHER person want?",
                     "Force yourself to see it from their side. If you can't name what they want, you're not ready. Write at least two distinct things they're probably optimizing for.",
                     lines=4)

    worksheet_prompt(doc, "04",
                     "What's my bottom line?",
                     "The number or answer I will NOT go below. Keep this in my head during the conversation. NEVER say it out loud.",
                     lines=2)

    worksheet_prompt(doc, "05",
                     "Which 3 tactics (of the 47) am I deploying?",
                     "Check your book. Not all 47. Pick 3 that fit this specific situation. One opener, one mid-conversation move, one exit move. Write the technique name and the page number.",
                     lines=5)

    worksheet_prompt(doc, "06",
                     "My opening line \u2014 word for word.",
                     "Draft the first sentence you will say when the conversation starts. The opening sets the frame for everything that follows. Do not wing this.",
                     lines=4)

    worksheet_prompt(doc, "07",
                     "The two objections/pushbacks I most expect.",
                     "For each one: (a) what they'll probably say, (b) my pre-drafted one-sentence response.",
                     lines=7)

    worksheet_prompt(doc, "08",
                     "My exit move.",
                     "How I end the conversation in a position of strength. What's the last sentence I say? What's the immediate next step I confirm before we separate?",
                     lines=4)

    worksheet_prompt(doc, "09",
                     "What's ONE thing I will NOT do in this conversation?",
                     "The specific trap I know I fall into. \u201cDon't defend my number.\u201d \u201cDon't apologize for bringing it up.\u201d \u201cDon't offer to think about it if they push back.\u201d Name your trap.",
                     lines=3)

    worksheet_prompt(doc, "10",
                     "Nerves check. Where am I on a 1-to-10 anxiety scale right now?",
                     "If you're above 6, run the 90-second nerves routine BEFORE you walk in. If you're below 3, you're probably under-preparing. Real preparation produces mild-to-moderate anxiety. That's fine.",
                     lines=2)

    # ==== PART 4: THE 3-TACTIC SELECTOR ====
    heading(doc, "THE 3-TACTIC SELECTOR", size=20, space_before=28)
    set_page_break_before(doc.paragraphs[-1])

    body(doc, "Question 5 of the worksheet asks you to pick 3 tactics from the 47 in the book. This table is the shortcut. Find your conversation type in the left column. Read across for the 3 tactics that fit best. Go deeper into the book only for the ones you chose.")

    selector = [
        ("Salary / raise conversation",
         "#5 The Consultant Stance, #11 The Strategic Pause, #41 The One-Line Email"),
        ("Hard conversation with partner",
         "#1 The Cold Open, #18 The Named Feeling, #17 The Earlier Agreement"),
        ("Reconnection with estranged family",
         "#1 The Cold Open, #6 The Confession First, #8 The Unrepayable Favor"),
        ("Custody / divorce mediation",
         "#4 The Assumption Audit, #19 The Pre-Understood Objection, #11 The Strategic Pause"),
        ("Hostile customer / client",
         "#4 The Assumption Audit, #18 The Named Feeling, #44 The Post-Agreement Bridge"),
        ("Contractor / service price dispute",
         "#11 The Strategic Pause, #13 The Impossible Answer, #41 The One-Line Email"),
        ("Hard conversation with a teenager or adult child",
         "#2 The Pre-Frame, #21 The Discovery Path, #37 The Unclosing"),
        ("Business / sales discovery call",
         "#33 The Reverse Qualifier, #21 The Discovery Path, #25 The Stakes Shift"),
        ("Confronting a manipulator",
         "#4 The Assumption Audit, #35 The Volcano, #40 The Public Witness"),
        ("Reconciling after a fight",
         "#6 The Confession First, #18 The Named Feeling, #41 The One-Line Email"),
        ("Job interview / being interviewed",
         "#5 The Consultant Stance, #22 The Misdirected Question, #33 The Reverse Qualifier"),
        ("Any conversation you're avoiding",
         "#2 The Pre-Frame, #11 The Strategic Pause, #43 The Minimum Viable Commitment"),
    ]
    for conv_type, tactics in selector:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.1)
        r1 = p.add_run(conv_type + "  \u2192  ")
        r1.font.name = BODY_FONT
        r1.font.size = Pt(10.5)
        r1.bold = True
        r1.font.color.rgb = DARK_RGB
        r2 = p.add_run(tactics)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK_RGB

    body(doc,
         "This table covers the twelve most common conversation types. For anything not on this list, grab the closest match and cross-reference the book's appendix, which indexes all 47 tactics by domain.",
         italic=True, color=BROWN_RGB, first_indent=0)

    # ==== PART 5: POST-CONVERSATION DEBRIEF ====
    heading(doc, "THE POST-CONVERSATION DEBRIEF", size=20, space_before=28)
    set_page_break_before(doc.paragraphs[-1])

    body(doc, "Fill this out within 2 hours of the conversation. This is the part most people skip. It's also the part that makes every future briefing better. The debrief is where you separate \u201cthe tactic worked\u201d from \u201cthe situation happened to go well.\u201d",
         italic=True, color=BROWN_RGB)

    worksheet_prompt(doc, "01", "What was the actual outcome?",
                     "Compare to question 2 on your briefing worksheet. Did you get what you wanted? Fully? Partially? Nothing?", lines=3)

    worksheet_prompt(doc, "02", "Which of the 3 tactics I picked did I actually deploy?",
                     "Be honest. Often you'll plan 3 and use only 1 or 2. That's information, not failure.", lines=4)

    worksheet_prompt(doc, "03", "What worked that I didn't plan?",
                     "Frequently the best move was spontaneous. Name it. That's a tactic you're developing that isn't formal yet.", lines=4)

    worksheet_prompt(doc, "04", "What didn't work that I thought would?",
                     "This is the single most valuable box on the debrief. Don't skip it.", lines=4)

    worksheet_prompt(doc, "05", "What's the one thing I'd change about my preparation next time?",
                     "Not \u201cme as a person.\u201d The preparation. This is how the next briefing gets sharper.", lines=4)

    # ==== PART 6: BLANK PRINT-AND-GO TEMPLATE ====
    heading(doc, "THE BRIEFING TEMPLATE (BLANK)", size=20, space_before=28)
    set_page_break_before(doc.paragraphs[-1])

    body(doc, "A clean, one-page version of the full worksheet. Print it and use it. Duplicate it as needed. Some users keep a stack in their desk drawer.", italic=True, color=BROWN_RGB)

    compact_prompts = [
        "Date / time / who:",
        "What I'm walking into (1 sentence):",
        "What I want (specific):",
        "What THEY want:",
        "My bottom line (don't speak this):",
        "My 3 tactics (+ page #):",
        "My opening line (word-for-word):",
        "Two objections I expect + my pre-drafted response:",
        "My exit move:",
        "The ONE thing I won't do:",
        "Nerves 1-10:",
    ]
    for c in compact_prompts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(c)
        r.font.name = TITLE_FONT
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = DARK_RGB
        ln = doc.add_paragraph()
        ln.paragraph_format.space_after = Pt(2)
        r = ln.add_run(" " + "_" * 90)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY_RGB
        ln = doc.add_paragraph()
        ln.paragraph_format.space_after = Pt(2)
        r = ln.add_run(" " + "_" * 90)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY_RGB

    # ==== CLOSING ====
    c = doc.add_paragraph()
    set_page_break_before(c)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(180)
    c.paragraph_format.space_after = Pt(12)
    r = c.add_run("FILL  //  REVIEW  //  WIN")
    r.font.name = TITLE_FONT
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    c2 = doc.add_paragraph()
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraph_format.space_after = Pt(18)
    r = c2.add_run("The best conversation is the one you walked into prepared,\nand walked out of with what you came for.")
    r.font.name = TITLE_FONT
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = BROWN_RGB

    c3 = doc.add_paragraph()
    c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c3.add_run("\u2014  Nate Harlan")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_RGB

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"Docx saved: {DOCX_OUT}")
    print(f"Size: {DOCX_OUT.stat().st_size} bytes")


if __name__ == "__main__":
    build_cover()
    build_docx()
